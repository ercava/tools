import docx

def debug_docx(path):
    doc = docx.Document(path)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if "pemahaman bacaan" in text:
            print(f"[{i}] {p.text.strip()}")
            # Print next few lines
            for j in range(1, 15):
                if i+j < len(doc.paragraphs):
                    n_text = doc.paragraphs[i+j].text.strip()
                    if n_text:
                        print(f"  + {n_text}")
            print("-" * 40)

if __name__ == "__main__":
    debug_docx("c:\\Users\\User\\Desktop\\utbk2025\\Projek MMA SNBT 2025 Lengkap.docx")
