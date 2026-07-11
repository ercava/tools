import docx

def preview_docx(path, num_paragraphs=150):
    try:
        doc = docx.Document(path)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        for i, p in enumerate(doc.paragraphs):
            if i > num_paragraphs: break
            text = p.text.strip()
            if text:
                print(f"[{i}] {text}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    preview_docx("c:\\Users\\User\\Desktop\\utbk2025\\Projek MMA SNBT 2025 Lengkap.docx")
